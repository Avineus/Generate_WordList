import enchant
import nltk
import re

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor

from docx.enum.text import WD_ALIGN_PARAGRAPH

from nltk.corpus import wordnet 

c = "bcdfghjklmnpqurstvwxyz"
v = "aeiou"

d = enchant.Dict("en_US")

def_str = ""
v_count = 2

document = Document()
document.add_heading('CVC Word List', 0)

p = document.add_paragraph('This book has list of CVC words with its associated part of speech. List of CVC words for each vowel is provided with list of ending words. Eg for a vowels list of all ab, ac, etc.. words. Table with CVC words respents NOUN as NN, VERB as VB, ADJECTIVE as ADJ, ADVERB as ADV.  Wordnet dictionary is used to get the meaning of each part of speech.')

document.add_page_break()

document.add_paragraph("Chapter 2: List of CVC words ", style='Heading 1')
document.add_paragraph("This chapter has list of all CVC words ")

### Table with list of cvc words

table_list = document.add_table(rows=50, cols=2)

column_count = 0
row_count = 0

word_count = 0

for i in v:
   for j in c:
      for k in c:
         word = k + i + j
         if d.check(word) : 

            if word_count == 100 :
               CVC_Liststr = "Next List of Words "
               document.add_paragraph(CVC_Liststr, style='Heading 1')

            if word_count != 0 and word_count % 100 == 0 :
               table_list = document.add_table(rows=50, cols=2)
               table_list.autofit = True
               table_list.style = 'Medium Grid 1 Accent 1'
               print ("Table Init at ", word_count)
               CVC_Liststr = "Next List of Words "
               document.add_paragraph(CVC_Liststr, style='Heading 1')
               print ("In Count:", word_count, word)


            #print ("Out Count:", word_count, word)
            cells = table_list.cell(row_count, column_count)
            table_list.autofit = True
            table_list.style = 'Medium Grid 1 Accent 1'

            text = k + i + j
            content = cells.add_paragraph(text, style='Heading 1')

            column_count = column_count + 1
            word_count = word_count + 1

            if column_count > 1 :
                column_count = 0
                row_count = row_count + 1

            if row_count == 50 :
               row_count = 0

print("Word Count is ", word_count)

document.add_page_break()
### Code to get dictionary meaning for CVC words in wordnet
### Loop to iterate for each letter in the vowel ####
### This loop will run 5 times a, e, i, o, u ###
for i in v:
   v_count = v_count + 1
   CVC_Vowel_str = "Chapter " + str(v_count) + ":" + "CVC words with wordnet meaning for " + i + " Vowel"
   document.add_paragraph(CVC_Vowel_str, style='Heading 1')
#   print(CVC_Vowel_str)

   ## Initialize Count of Words to 0
   count = 0

   ### This loop is to iterate to get C after vowels
   #### In this case it is a[a-z], e[a-z], etc..
   ### This loop will iterate for 21 consonants
   for j in c:
      ### This loop is to iterate to get C at the start
      ### In this case it is eg) at, am will be iterated to get first letter
      ### [a-z]at, [a-z]am, etc..
      ### This loop will iterate for 21 consonants

      count = 0
      for k in c:
         ### Compare formed 3 letter word to the dictionary
         ## Only if word is available in dictory then print
         word = k + i + j

         if d.check(word) : 

            n_len = len(wordnet.synsets(word, pos='n'))
            v_len = len(wordnet.synsets(word, pos='v'))
            a_len = len(wordnet.synsets(word, pos='a'))
            r_len = len(wordnet.synsets(word, pos='r'))

            if (n_len + v_len + a_len + r_len) == 0 :
               break

            if count == 0 :
               CVC_end_str = "Words ending with " + i + j
               document.add_paragraph(CVC_end_str, style='Heading 2')

               table = document.add_table(rows=1, cols=2)

               table.autofit = True
               table.style = 'Medium Grid 1 Accent 1'

               hdr_cells = table.rows[0].cells
               hdr_cells[0].text = 'CVC Words'
               hdr_cells[1].text = 'Parts of Speech'

            count = count + 1;

            row_cells = table.add_row().cells
            row_cells[0].text = str(word)

            #pos_str = "NN=" + str(n_len) + " VB=" + str(v_len) + " ADJ=" + str(a_len) + " ADV=" + str(r_len) 
            #print(word + " " +  pos_str)
	    #row_cells[1].text = str(pos_str)

            if n_len != 0 :
               def_str = "NOUNS \n"

            ind = 0
            for noun in wordnet.synsets(word, pos='n'):
                ind = ind + 1
                noun_str = noun.definition()
                def_str = def_str + "[" + str(ind) + "]" + noun_str + "\n"

            if v_len != 0:
               def_str = def_str + "\n VERBS \n"

            ind = 0
            for verb in wordnet.synsets(word, pos='v'):
                ind = ind + 1
                verb_str = verb.definition()
                def_str = def_str + "[" + str(ind) + "]" + verb_str + "\n"

            if a_len != 0:
                def_str = def_str + "\n ADJECTIVES \n"

            ind = 0
            for adjective in wordnet.synsets(word, pos='a'):
                ind = ind + 1
                adjective_str = adjective.definition()
                def_str = def_str + "[" + str(ind) + "]" + adjective_str + "\n"

            if r_len != 0:
               def_str = def_str + "\n ADVERBS \n"

            ind = 0
            for adverb in wordnet.synsets(word, pos='r'):
                ind = ind + 1
                adverb_str = adverb.definition()
                def_str = def_str + "[" + str(ind) + "]" + adverb_str + "\n"

            row_cells[1].text = str(def_str)
            def_str = ""

   document.add_page_break()

   document.save('Output/CVC_Word.docx')
            
            

