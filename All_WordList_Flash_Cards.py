import enchant
import nltk
import re

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Cm

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT

from nltk.corpus import wordnet 

c = "bcdfghjklmnpqurstvwxyz"
v = "aeiou"

d = enchant.Dict("en_US")

def_str = ""
v_count = 2

document = Document()
document.add_heading('Need for Word List', 0)

p = document.add_paragraph('This book has list of CVC, CCVC, CVCC, CVVC and CCVCC words flash cards. Each word list is separated as 200 rows in a table. This will be helpful for parents to introduce more words with phonics. This list is generated by combining vowels and consonants and comparing with online dictionary to find meaning.')

document.add_page_break()

document.add_paragraph("Chapter 2: List of CVC words ", style='Heading 1')
document.add_paragraph("This chapter has list of all CVC words ")

### Table with list of cvc words

table_list = document.add_table(rows=50, cols=4)

column_count = 0
row_count = 0

word_count = 0

for i in v:
   for j in c:
      for k in c:
         word = k + i + j
         if d.check(word) : 

            if word_count == 200 :
               CVC_Liststr = "Next List of Words "
               #document.add_paragraph(CVC_Liststr, style='Heading 1')

            if word_count != 0 and word_count % 200 == 0 :
               table_list = document.add_table(rows=50, cols=4)
               table_list.autofit = True
               table_list.style = 'Medium Grid 1 Accent 1'
               print ("Table Init at ", word_count)
               CVC_Liststr = "Next List of Words "
              # document.add_paragraph(CVC_Liststr, style='Heading 1')
               print ("In Count:", word_count, word)


            #print ("Out Count:", word_count, word)
            cells = table_list.cell(row_count, column_count)
           # table_list.rows.height_rule = WD_ROW_HEIGHT.EXACTLY
            #table_list.rows.height = Inches(1.9655)
            table_list.autofit = True
           # table_list.cell(row_count, column_count).width = Inches(1.9655)
            table_list.style = 'Medium Grid 1 Accent 1'

            text = k + i + j 
            content = cells.add_paragraph("\n\n" + text + "\n\n", style='Heading 1')

            column_count = column_count + 1
            word_count = word_count + 1

            if column_count > 3 :
                column_count = 0
                row_count = row_count + 1

            if row_count == 50 :
               row_count = 0

print("Word Count is ", word_count)

document.add_page_break()

document.add_paragraph("Chapter 3: List of CCVC words ", style='Heading 1')
document.add_paragraph("This chapter has list of all CCVC words ")

c = "bcdfghjklmnpqurstvwxyz"
v = "aeiou"

d = enchant.Dict("en_US")

def_str = ""
v_count = 2

### Table with list of ccvc words

table_list = document.add_table(rows=50, cols=4)
column_count = 0
row_count = 0
word_count = 0

for i in v:
   for j in c:
      for k in c:
         for l in c:
           word = l + k + i + j
           if d.check(word) : 

              text = l + k + i + j
              n_len = len(wordnet.synsets(word, pos='n'))
              v_len = len(wordnet.synsets(word, pos='v'))
              a_len = len(wordnet.synsets(word, pos='a'))
              r_len = len(wordnet.synsets(word, pos='r'))

              if (n_len + v_len + a_len + r_len) == 0 :
               break


              if word_count == 200 :
                CCVC_Liststr = "Next List of Words "
                #document.add_paragraph(CCVC_Liststr, style='Heading 1')

              if word_count != 0 and word_count % 200 == 0 :
                 table_list = document.add_table(rows=50, cols=4)
                 table_list.autofit = True
                 table_list.style = 'Medium Grid 1 Accent 3'
                 print ("Table Init at ", word_count)
                 CCVC_Liststr = "Next List of Words "
                # document.add_paragraph(CCVC_Liststr, style='Heading 1')
                 print ("In Count:", word_count, word)


              #print ("Out Count:", word_count, word)
              cells = table_list.cell(row_count, column_count)
              table_list.autofit = True
              table_list.style = 'Medium Grid 1 Accent 3'

              text =  l + k + i + j
              content = cells.add_paragraph("\n\n" + text + "\n\n", style='Heading 1')

              column_count = column_count + 1
              word_count = word_count + 1

              if column_count > 3 :
                 column_count = 0
                 row_count = row_count + 1

              if row_count == 50 :
                 row_count = 0

document.add_page_break()

document.add_paragraph("Chapter 4: List of CVCC words ", style='Heading 1')
document.add_paragraph("This chapter has list of all CVCC words ")

c = "bcdfghjklmnpqurstvwxyz"
v = "aeiou"

d = enchant.Dict("en_US")

def_str = ""
v_count = 2

### Table with list of cvcc words

table_list = document.add_table(rows=50, cols=4)
column_count = 0
row_count = 0
word_count = 0

for i in v:
   for j in c:
      for k in c:
         for l in c:
           word = k + i + j + l
           if d.check(word) : 

              text = k + i + j + l
              n_len = len(wordnet.synsets(word, pos='n'))
              v_len = len(wordnet.synsets(word, pos='v'))
              a_len = len(wordnet.synsets(word, pos='a'))
              r_len = len(wordnet.synsets(word, pos='r'))

              if (n_len + v_len + a_len + r_len) == 0 :
               break


              if word_count == 200 :
                CVCC_Liststr = "Next List of Words "
                #document.add_paragraph(CVCC_Liststr, style='Heading 1')

              if word_count != 0 and word_count % 200 == 0 :
                 table_list = document.add_table(rows=50, cols=4)
                 table_list.autofit = True
                 table_list.style = 'Medium Grid 1 Accent 2'
                 print ("Table Init at ", word_count)
                 CVCC_Liststr = "Next List of Words "
                 #document.add_paragraph(CVCC_Liststr, style='Heading 1')
                 print ("In Count:", word_count, word)


              #print ("Out Count:", word_count, word)
              cells = table_list.cell(row_count, column_count)
              table_list.autofit = True
              table_list.style = 'Medium Grid 1 Accent 2'

              text = k + i + j + l
              content = cells.add_paragraph("\n\n" + text + "\n\n", style='Heading 1')

              column_count = column_count + 1
              word_count = word_count + 1

              if column_count > 3 :
                 column_count = 0
                 row_count = row_count + 1

              if row_count == 50 :
                 row_count = 0

document.add_page_break()

document.add_paragraph("Chapter 5: List of CVVC words ", style='Heading 1')
document.add_paragraph("This chapter has list of all CVVC words ")

c = "bcdfghjklmnpqurstvwxyz"
v = "aeiou"

d = enchant.Dict("en_US")

def_str = ""
v_count = 2

### Table with list of cvvc words

table_list = document.add_table(rows=50, cols=4)
column_count = 0
row_count = 0
word_count = 0

for i in v:
   for j in c:
      for k in c:
           word = j + i + i + k
           if d.check(word) : 

              text = j + i + i + k
              n_len = len(wordnet.synsets(word, pos='n'))
              v_len = len(wordnet.synsets(word, pos='v'))
              a_len = len(wordnet.synsets(word, pos='a'))
              r_len = len(wordnet.synsets(word, pos='r'))

              if (n_len + v_len + a_len + r_len) == 0 :
               break


              if word_count == 200 :
                CVVC_Liststr = "Next List of Words "
                #document.add_paragraph(CVVC_Liststr, style='Heading 1')

              if word_count != 0 and word_count % 200 == 0 :
                 table_list = document.add_table(rows=50, cols=4)
                 table_list.autofit = True
                 table_list.style = 'Medium Grid 1 Accent 3'
                 print ("Table Init at ", word_count)
                 CVVC_Liststr = "Next List of Words "
                 #document.add_paragraph(CVVC_Liststr, style='Heading 1')
                 print ("In Count:", word_count, word)


              #print ("Out Count:", word_count, word)
              cells = table_list.cell(row_count, column_count)
              table_list.autofit = True
              table_list.style = 'Medium Grid 1 Accent 5'

              text =  j + i + i + k 
              content = cells.add_paragraph("\n\n" + text + "\n\n", style='Heading 1')

              column_count = column_count + 1
              word_count = word_count + 1

              if column_count > 3 :
                 column_count = 0
                 row_count = row_count + 1

              if row_count == 50 :
                 row_count = 0

document.add_page_break()

document.add_paragraph("Chapter 6: List of CCVCC words ", style='Heading 1')
document.add_paragraph("This chapter has list of all CCVCC words ")

### Table with list of ccvcc words

table_list = document.add_table(rows=50, cols=4)
column_count = 0
row_count = 0
word_count = 0

for i in v:
   for j in c:
      for k in c:
         for l in c:
           for m in c:
             word = l + k + i + j + m
             if d.check(word) : 

                text = l + k + i + j + m
                n_len = len(wordnet.synsets(word, pos='n'))
                v_len = len(wordnet.synsets(word, pos='v'))
                a_len = len(wordnet.synsets(word, pos='a'))
                r_len = len(wordnet.synsets(word, pos='r'))

                if (n_len + v_len + a_len + r_len) == 0 :
                   break


                if word_count == 200 :
                   CCVCC_Liststr = "Next List of Words "
                  # document.add_paragraph(CCVCC_Liststr, style='Heading 1')

                if word_count != 0 and word_count % 200 == 0 :
                   table_list = document.add_table(rows=50, cols=4)
                   table_list.autofit = True
                   table_list.style = 'Medium Grid 1 Accent 4'
                   print ("Table Init at ", word_count)
                   CCVCC_Liststr = "Next List of Words "
                  # document.add_paragraph(CCVCC_Liststr, style='Heading 1')
                   print ("In Count:", word_count, word)


                #print ("Out Count:", word_count, word)
                cells = table_list.cell(row_count, column_count)
                table_list.autofit = True
                table_list.style = 'Medium Grid 1 Accent 4'

                text =  l + k + i + j + m
                content = cells.add_paragraph("\n\n" + text + "\n\n", style='Heading 1')

                column_count = column_count + 1
                word_count = word_count + 1

                if column_count > 3 :
                   column_count = 0
                   row_count = row_count + 1

                if row_count == 50 :
                   row_count = 0

document.save('Output/All_WordList_Flash.docx')
            
            

