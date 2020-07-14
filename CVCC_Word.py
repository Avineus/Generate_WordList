import enchant
import nltk
import re

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor

from docx.enum.text import WD_ALIGN_PARAGRAPH

from nltk.corpus import wordnet 

c = "bcdfghjklmnpqurstvwxyz"
v = "aeiou"

d = enchant.Dict("en_US")

def_str = ""
v_count = 2

document = Document()
document.add_heading('Need for CVCC Word List', 0)

p = document.add_paragraph('CVCC words are useful for kids to start early reading. There are 5 vowels which can be combined with consonants to form words. Words which are formed with consonants +  vowels + consonants + consonants are called CVCC words.Blending words are essential for early readers, this will become easy with familiarity of CVCC words. In this book there are about 1043 list of all possible CVCC words. This book has collection of CVCC words starting with  ba, pa, ra, be, ti, bo, bu, etc.. This book has list of CCVC words with its associated part of speech. Table with CVCC words are represented as NOUN, VERB, ADJECTIVE and ADVERB . Flash Card for 1043 flash cards is also available')

document.add_page_break()

### Code to get dictionary meaning for CVC words in wordnet
### Loop to iterate for each letter in the vowel ####
### This loop will run 5 times a, e, i, o, u ###
for i in v:
   v_count = v_count + 1
   CVCC_Vowel_str = "2 " + "CVCC words Dictionary for vowel " + i 
   document.add_paragraph(CVCC_Vowel_str, style='Heading 1')

   ## Initialize Count of Words to 0
   count = 0

   ### This loop is to iterate to get C after vowels
   #### In this case it is a[a-z], e[a-z], etc..
   ### This loop will iterate for 21 consonants
   for j in c:
      ### This loop is to iterate to get C at the start
      ### In this case it is eg) at, am will be iterated to get first letter
      ### [a-z]at, [a-z]am, etc..
      ### This loop will iterate for 21 consonants

      count = 0
      for k in c:
         for l in c:
           ### Compare formed 4 letter word to the dictionary
           ## Only if word is available in dictory then check in wordnet 
           word = k + i + j + l

           if d.check(word) : 

             n_len = len(wordnet.synsets(word, pos='n'))
             v_len = len(wordnet.synsets(word, pos='v'))
             a_len = len(wordnet.synsets(word, pos='a'))
             r_len = len(wordnet.synsets(word, pos='r'))

             if (n_len + v_len + a_len + r_len) == 0 :
                break

             if count == 0 :
                CVCC_start_str = "Words Starting with " + k + i + j
                document.add_paragraph(CVCC_start_str, style='Heading 2')

                table = document.add_table(rows=1, cols=2)

                table.autofit = True
                table.style = 'Medium Grid 1 Accent 2'

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'CVCC Words'
                hdr_cells[1].text = 'Parts of Speech'

             count = count + 1;

             row_cells = table.add_row().cells
             row_cells[0].text = str(word)

             if n_len != 0 :
               def_str = "NOUNS \n"

             ind = 0
             for noun in wordnet.synsets(word, pos='n'):
                ind = ind + 1
                noun_str = noun.definition()
                def_str = def_str + "[" + str(ind) + "]" + noun_str + "\n"

             if v_len != 0:
               def_str = def_str + "\n VERBS \n"

             ind = 0
             for verb in wordnet.synsets(word, pos='v'):
                ind = ind + 1
                verb_str = verb.definition()
                def_str = def_str + "[" + str(ind) + "]" + verb_str + "\n"

             if a_len != 0:
                def_str = def_str + "\n ADJECTIVES \n"

             ind = 0
             for adjective in wordnet.synsets(word, pos='a'):
                ind = ind + 1
                adjective_str = adjective.definition()
                def_str = def_str + "[" + str(ind) + "]" + adjective_str + "\n"

             if r_len != 0:
               def_str = def_str + "\n ADVERBS \n"

             ind = 0
             for adverb in wordnet.synsets(word, pos='r'):
                ind = ind + 1
                adverb_str = adverb.definition()
                def_str = def_str + "[" + str(ind) + "]" + adverb_str + "\n"

             row_cells[1].text = str(def_str)
             def_str = ""

   document.add_page_break()


   CVCC_Flash_str = "7 Flash Cards for 1043 CVCC Words" 
   document.add_paragraph(CVCC_Flash_str, style='Heading 1')

### Table with list of cvc words

table_list = document.add_table(rows=50, cols=4)
column_count = 0
row_count = 0
word_count = 0

for i in v:
   for j in c:
      for k in c:
         for l in c:
           word = k + i + j + l
           if d.check(word) : 

              text = k + i + j + l
              n_len = len(wordnet.synsets(word, pos='n'))
              v_len = len(wordnet.synsets(word, pos='v'))
              a_len = len(wordnet.synsets(word, pos='a'))
              r_len = len(wordnet.synsets(word, pos='r'))

              if (n_len + v_len + a_len + r_len) == 0 :
               break


              if word_count == 200 :
                CVCC_Liststr = "Next List of Words "
                document.add_paragraph(CVCC_Liststr, style='Heading 1')

              if word_count != 0 and word_count % 200 == 0 :
                 table_list = document.add_table(rows=50, cols=4)
                 table_list.autofit = True
                 table_list.style = 'Medium Grid 1 Accent 2'
                 print ("Table Init at ", word_count)
                 CVCC_Liststr = "Next List of Words "
                 document.add_paragraph(CVCC_Liststr, style='Heading 1')
                 print ("In Count:", word_count, word)


              #print ("Out Count:", word_count, word)
              cells = table_list.cell(row_count, column_count)
              table_list.autofit = True
              table_list.style = 'Medium Grid 1 Accent 2'

              text = k + i + j + l
              content = cells.add_paragraph(text, style='Heading 1')

              column_count = column_count + 1
              word_count = word_count + 1

              if column_count > 3 :
                 column_count = 0
                 row_count = row_count + 1

              if row_count == 50 :
                 row_count = 0

              document.save('Output/CVCC_Word.docx')

print("Word Count is ", word_count)


